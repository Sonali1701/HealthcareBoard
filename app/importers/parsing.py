"""Best-effort text extraction + field parsing for healthcare resumes.

This is heuristic, not AI — it reliably pulls email/phone and detects
healthcare specialties, professions, certifications and US states by keyword.
Imported profiles are flagged ``source=resume_parse`` so a human can review and
correct them in-app afterward.
"""
from __future__ import annotations

import io
import re
from pathlib import Path

# --- Domain vocabularies --------------------------------------------------

PROFESSIONS = {
    "RN": ["registered nurse", " rn ", " rn,", "rn ", "bsn", "registered nurse (rn)"],
    "LPN": ["licensed practical nurse", " lpn", "lvn"],
    "CNA": ["certified nursing assistant", " cna"],
    "NP": ["nurse practitioner", " np ", "fnp", "acnp", "pmhnp"],
    "CRNA": ["nurse anesthetist", "crna"],
    "CNM": ["certified nurse midwife", "nurse midwife", " cnm"],
    "MD": ["physician", " md ", " m.d", "doctor of medicine", "mbchb"],
    "PA": ["physician assistant", "pa-c"],
    "RT": ["respiratory therapist", "registered respiratory"],
    "PT": ["physical therapist", " dpt", "physical therapy"],
    "OT": ["occupational therapist", "occupational therapy"],
    "PharmD": ["pharmacist", "pharm.d", "pharmd"],
}

SPECIALTIES = {
    "ICU": ["icu", "intensive care", "critical care", "ccu", "sicu", "micu"],
    "ER": ["emergency", " er ", " ed ", "emergency department", "emergency room"],
    "OR": ["operating room", "perioperative", " or nurse", "surgical services"],
    "PICU": ["picu", "pediatric intensive"],
    "NICU": ["nicu", "neonatal"],
    "Labor & Delivery": ["labor and delivery", "labor & delivery", "l&d", "postpartum",
                         "antepartum", "mother baby"],
    "Med-Surg": ["med surg", "med-surg", "medical surgical", "medical-surgical"],
    "Telemetry": ["telemetry", "tele unit", "step down", "stepdown", "pcu"],
    "Oncology": ["oncology", "hematology", "chemo"],
    "Cath Lab": ["cath lab", "cardiac cath", "interventional"],
    "PACU": ["pacu", "post anesthesia", "recovery room"],
    "Dialysis": ["dialysis", "nephrology", "hemodialysis"],
    "Home Health": ["home health", "home care"],
    "Psych": ["psychiatric", "behavioral health", "psych "],
}

CERTIFICATIONS = ["BLS", "ACLS", "PALS", "CCRN", "TNCC", "NRP", "CEN", "ATLS",
                  "CNOR", "RNC", "CPN", "CMSRN", "PCCN", "TCRN", "AWHONN",
                  "STABLE", "NIHSS", "EKG"]

US_STATES = {
    "AL": "alabama", "AK": "alaska", "AZ": "arizona", "AR": "arkansas",
    "CA": "california", "CO": "colorado", "CT": "connecticut", "DE": "delaware",
    "FL": "florida", "GA": "georgia", "HI": "hawaii", "ID": "idaho",
    "IL": "illinois", "IN": "indiana", "IA": "iowa", "KS": "kansas",
    "KY": "kentucky", "LA": "louisiana", "ME": "maine", "MD": "maryland",
    "MA": "massachusetts", "MI": "michigan", "MN": "minnesota", "MS": "mississippi",
    "MO": "missouri", "MT": "montana", "NE": "nebraska", "NV": "nevada",
    "NH": "new hampshire", "NJ": "new jersey", "NM": "new mexico", "NY": "new york",
    "NC": "north carolina", "ND": "north dakota", "OH": "ohio", "OK": "oklahoma",
    "OR": "oregon", "PA": "pennsylvania", "RI": "rhode island", "SC": "south carolina",
    "SD": "south dakota", "TN": "tennessee", "TX": "texas", "UT": "utah",
    "VT": "vermont", "VA": "virginia", "WA": "washington", "WV": "west virginia",
    "WI": "wisconsin", "WY": "wyoming",
}

# Post-nominal credentials to strip from a name line / detect profession from.
CREDENTIALS = [
    "MD", "DO", "MBBS", "MBChB", "DMD", "DDS", "DPM", "DPT", "PharmD", "PsyD",
    "PhD", "MPH", "MSc", "MS", "MBA", "MHA", "MFA", "BFA", "BA", "BS",
    "RN", "BSN", "MSN", "NP",
    "FNP", "DNP", "PA", "PA-C", "FACP", "FAAP", "FACS", "FACC", "FAAD",
    "FACRO", "FASN", "FCAAAI", "FCAAI",
    "FAAAAI", "FACAAI", "MSCR", "FACE", "I", "II", "III", "IV",
]
PROFESSION_FROM_CRED = ["DO", "MD", "MBBS", "MBChB", "DPM", "DMD", "DDS", "PharmD",
                        "DPT", "DNP", "NP", "PA", "RN"]

# --- Provider categories (Physicians / Nursing / Allied / APP / Others) ---
# Used by the Providers directory. APP is checked first because CRNA/NP names
# contain the word "nurse" and would otherwise fall into Nursing.
PROVIDER_CATEGORIES = ["Physicians", "Nursing", "Allied", "APP", "Others"]

_NURSING_SPECIALTY_LABELS = {
    "icu", "er", "picu", "nicu", "labor & delivery", "med-surg",
    "telemetry", "oncology", "pacu", "or", "operating room", "dialysis",
}

_APP_CODES = {"NP", "CRNA", "FNP", "FNP-C", "FNP-BC", "AGNP", "PMHNP",
              "PMHNP-BC", "ACNP", "AGACNP", "PNP", "WHNP", "NP-C"}
_APP_KW = ["nurse practitioner", "nurse anesthetist", "crna",
           "certified registered nurse anesthetist"]
_PHYS_CODES = {"MD"}
_PHYS_KW = ["medical doctor", "doctor of medicine", "family medicine"]
_NURSE_CODES = {"RN", "LPN", "CNA"}
_NURSE_KW = ["registered nurse", "licensed practical nurse",
             "certified nursing assistant"]
_ALLIED_CODES = {"RAD", "RADTECH", "RTR", "ARRT", "RDMS", "RVT", "RCIS",
                 "CNMT", "RDCS", "CT", "MRI"}
_ALLIED_KW = [
    "radiologic technologist", "rad tech", "radiographer", "x-ray", "x ray",
    "ct technologist", "ct tech", "mri technologist", "mri tech",
    "mammography technologist", "mammography tech", "mammographer",
    "ultrasound technologist", "ultrasound tech", "sonographer",
    "echocardiography technologist", "echo tech", "echo technologist",
    "vascular technologist", "vascular tech", "nuclear medicine technologist",
    "nuclear medicine tech", "interventional radiology technologist",
    "interventional radiology tech", "ir technologist", "ir tech",
    "cardiac cath lab technologist", "cardiac cath lab tech",
    "cath lab technologist", "cath lab tech", "radiology technologist",
]

ABPS = "american board of physician specialties"

# License/title code -> full name. Curated to the licenses that matter for a
# healthcare staffing board (used for the filter dropdown labels AND to make the
# full name searchable, e.g. searching "registered nurse" finds RN profiles).
LICENSE_FULL_NAMES = {
    "RN": "Registered Nurse",
    "LPN": "Licensed Practical Nurse",
    "LVN": "Licensed Vocational Nurse",
    "CNA": "Certified Nursing Assistant",
    "NP": "Nurse Practitioner",
    "FNP": "Family Nurse Practitioner",
    "DNP": "Doctor of Nursing Practice",
    "CRNA": "Certified Registered Nurse Anesthetist",
    "CNM": "Certified Nurse Midwife",
    "PA": "Physician Assistant",
    "MD": "Physician",
    "DO": "Doctor of Osteopathic Medicine",
    "RT": "Respiratory Therapist",
    "PT": "Physical Therapist",
    "OT": "Occupational Therapist",
}

# Words that appear in a bad "name" but essentially never in a real person's
# name — roles, credentials, résumé/section words, and generic filler. We detect
# junk by STRUCTURE (any of these as a word, a run-on, or a digit) rather than by
# listing specific bad names, so new variants are caught automatically.
JUNK_NAME_WORDS = {
    # placeholders
    "unknown", "candidate", "provider", "providers", "resume", "cv", "n/a", "na",
    "none", "member", "applicant", "profile", "portfolio",
    # roles / titles
    "registered", "certified", "licensed", "nurse", "nursing", "physician",
    "surgeon", "doctor", "technician", "technologist", "practitioner",
    "assistant", "associate", "provider", "professional", "administrator",
    "administrative", "coordinator", "specialist", "therapist", "pharmacist",
    "director", "manager", "management", "supervisor", "consultant", "clinician",
    "caregiver", "aide", "worker", "staff",
    # résumé / section words
    "curriculum", "vitae", "objective", "summary", "references", "reference",
    "experience", "experienced", "qualifications", "education", "skills", "skill",
    "certifications", "licensure", "employment", "history",
    # generic filler
    "healthcare", "medical", "clinical", "hospital", "university", "college",
    "career", "seeking", "dedicated", "motivated", "organized", "acquired",
    "regional", "center", "travel", "staffing", "solutions", "services",
    "department", "unit", "team", "group",
    # credentials that leak in as a name token
    "adn", "bsn", "msn", "dnp", "aprn", "faan", "mba", "mph", "phd",
    # placeholders the importers emit when they cannot find a name
    "not", "found", "nil", "blank", "empty", "tbd", "missing", "undefined",
    # résumé section headings that get read as the name line
    "personal", "information", "contact", "details", "address", "areas",
    "expertise", "competencies", "core", "key", "highlights", "achievements",
    "accomplishments", "responsibilities", "projects", "activities", "interests",
    "additional", "background", "overview", "introduction", "declaration",
    "hobbies", "languages", "awards", "publications", "training", "courses",
    "workshops", "affiliations", "memberships", "strengths", "attributes",
    # job-board scrape artifacts
    "view", "lead", "leads", "apply", "click", "page", "search", "results",
    # non-clinical role words that show up as names on IT résumés
    "developer", "engineer", "analyst", "architect", "programmer", "designer",
    "tester", "scrum", "sr", "jr", "senior", "junior", "intern", "freelance",
    # Function words of 3+ letters only. Two-letter words are excluded on
    # purpose: "An", "To", "Bo", "Le" are real Vietnamese names and rejecting
    # them would erase real people.
    "and", "the", "for", "with", "from", "more", "than", "are", "was", "were",
    "but", "not",
}

# A surname that is really a US state code means the parser grabbed a location
# line ("San TX", "Columbus OH") instead of a person.
_STATE_CODES = {
    "AL", "AK", "AZ", "AR", "CA", "CO", "CT", "DE", "FL", "GA", "HI", "ID",
    "IL", "IN", "IA", "KS", "KY", "LA", "ME", "MD", "MA", "MI", "MN", "MS",
    "MO", "MT", "NE", "NV", "NH", "NJ", "NM", "NY", "NC", "ND", "OH", "OK",
    "OR", "PA", "RI", "SC", "SD", "TN", "TX", "UT", "VT", "VA", "WA", "WV",
    "WI", "WY", "DC",
}
NAME_PLACEHOLDERS = JUNK_NAME_WORDS   # back-compat alias
_JUNK_SUBSTRINGS = tuple(w for w in JUNK_NAME_WORDS if len(w) >= 6)


def is_real_name(first, last) -> bool:
    """True only if (first, last) looks like an actual person's name.

    Structural, not a blocklist: a name is junk if it contains a digit, a role /
    résumé word, or a run-on concatenation (e.g. 'Registerednursefor',
    'Organizedandprofessional') — so we don't have to enumerate bad names.
    """
    f = (first or "").strip().lower().strip(".")
    if not f or not any(c.isalpha() for c in f):
        return False
    lastname = (last or "").strip().strip(".")
    if lastname.upper() in _STATE_CODES:
        return False                        # "San TX" — a location, not a person
    full = f"{f} {lastname.lower()}".strip()
    if any(ch.isdigit() for ch in full):
        return False
    for w in re.split(r"[\s\-]+", full):
        w = w.strip(".")
        if not w:
            continue
        if w in JUNK_NAME_WORDS:            # a whole word is a role/résumé word
            return False
        if len(w) >= 16:                    # implausibly long single token = run-on
            return False
        if len(w) >= 12 and any(sub in w for sub in _JUNK_SUBSTRINGS):
            return False                    # long run-on that embeds a junk word
    return True


def classify_provider(profession_type=None, specialty=None,
                      headline=None, title=None) -> str:
    """Bucket a profile into the four requested groups, or ``Others``."""
    code = (profession_type or "").upper().strip().strip(".")
    text = " ".join(x for x in (profession_type, specialty, headline, title) if x).lower()
    if code in _APP_CODES or any(k in text for k in _APP_KW):
        return "APP"
    if code in _PHYS_CODES or any(k in text for k in _PHYS_KW) \
            or re.search(r"\bm\.?d\.?\b", text):
        return "Physicians"
    if code in _NURSE_CODES or any(k in text for k in _NURSE_KW):
        return "Nursing"
    if code in _ALLIED_CODES or any(k in text for k in _ALLIED_KW):
        return "Allied"
    specialty_label = (specialty or "").strip().lower()
    if specialty_label in _NURSING_SPECIALTY_LABELS and code in {"", "NURSE"}:
        return "Nursing"
    return "Others"


def primary_american_board(cert_names) -> str | None:
    """First 'American Board of …' certification, excluding ABPS."""
    for c in cert_names or []:
        cl = str(c).strip().lower()
        if cl.startswith("american board of") and cl != ABPS:
            return str(c).strip()
    return None


EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}")
YEARS_RE = re.compile(r"(\d{1,2})\+?\s*(?:years|yrs?)\b", re.IGNORECASE)
NPI_RE = re.compile(r"\bNPI[:#\s]*([0-9]{10})\b", re.IGNORECASE)
STATE_CODE_RE = "|".join(US_STATES)
STATE_NAME_RE = "|".join(re.escape(name) for name in US_STATES.values())
CITY_STATE_RE = re.compile(
    rf"\b([A-Z][A-Za-z.' -]{{1,45}}?),\s*({STATE_CODE_RE})\b(?:\s+\d{{5}}(?:-\d{{4}})?)?"
)
CITY_FULL_STATE_RE = re.compile(
    rf"\b([A-Z][A-Za-z.' -]{{1,45}}?),\s*({STATE_NAME_RE})\b(?:\s+\d{{5}}(?:-\d{{4}})?)?",
    re.IGNORECASE,
)
SECTION_TEXT = {
    "summary", "objective", "profile", "membership", "memberships",
    "organizational", "education", "education & training", "licensure",
    "licensure & certifications", "certification", "certifications",
    "certifications & licensure", "licenses", "skills", "experience",
    "clinical experience", "professional experience", "healthcare experience",
    "work experience", "employment", "references",
}
NOISY_WORDS = {
    "summary", "membership", "organizational", "education", "licensure",
    "certification", "certifications", "experience", "social media",
    "proficiency", "training", "provider training", "surgical", "clinical",
}
ADDRESS_WORDS = {
    "road", "rd", "street", "st", "drive", "dr", "court", "ct", "way",
    "avenue", "ave", "blvd", "boulevard", "lane", "ln", "apt", "suite",
    "floor", "unit", "po box", "p.o. box",
}
CITY_BAD_WORDS = SECTION_TEXT | {
    "provider", "hospital", "clinic", "department", "surgical", "summary",
    "education", "experience", "certification", "certifications", "licensure",
    "congress", "conference", "academy", "association", "society", "university",
    "college", "dermatology", "pediatric", "medical", "region",
}


# --- Text extraction ------------------------------------------------------

def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise ValueError(f"Unsupported resume type: {suffix} (use .pdf or .docx)")


# A "glued" token: a long word with an internal lower->upper transition, e.g.
# "AssistedLiving" / "SpecialtyInformation". These appear when the PDF has no
# space glyphs and pypdf's default mode runs words together.
_GLUED_TOKEN = re.compile(r"\b(?=\w{12,}\b)[A-Za-z]*[a-z][A-Z][A-Za-z]*\b")
_ALPHA_WORD = re.compile(r"[A-Za-z]{2,}")


def _glue_score(text: str) -> float:
    """Fraction of words that look run-together. Lower is better."""
    words = _ALPHA_WORD.findall(text or "")
    if not words:
        return 1.0
    return len(_GLUED_TOKEN.findall(text)) / len(words)


def _pdf_text(reader) -> str:
    """PDF text, preferring whichever pypdf mode preserves word spacing.

    pypdf's default mode silently drops inter-word spaces on some documents
    ("Unit:LTC,Skilled,AssistedLiving"), which also hides section headings from
    the résumé parser. Layout mode keeps the spacing but can interleave columns
    on multi-column résumés, so we extract both and keep the cleaner one rather
    than committing to either.
    """
    def _render(mode: str | None) -> str:
        out = []
        for page in reader.pages:
            try:
                out.append((page.extract_text(extraction_mode=mode) if mode
                            else page.extract_text()) or "")
            except Exception:
                out.append("")
        return "\n".join(out)

    plain = _render(None)
    try:
        layout = _render("layout")
    except Exception:
        return plain
    if not _ALPHA_WORD.search(layout):
        return plain
    if not _ALPHA_WORD.search(plain):
        return layout
    # Prefer layout unless it is measurably worse at keeping words apart.
    return layout if _glue_score(layout) <= _glue_score(plain) else plain


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    return _pdf_text(PdfReader(str(path)))


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" ".join(c.text for c in row.cells))
    return "\n".join(lines)


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    suffix = Path(filename or "resume").suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        return _pdf_text(PdfReader(io.BytesIO(data)))
    if suffix == ".docx":
        from docx import Document

        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                lines.append(" ".join(c.text for c in row.cells))
        return "\n".join(lines)
    raise ValueError(f"Unsupported resume type: {suffix}")


# --- Field parsing --------------------------------------------------------

def _name_from_filename(path: Path) -> tuple[str, str]:
    stem = re.sub(r"(?i)(resume|cv|_|-)", " ", path.stem).strip()
    parts = [p for p in stem.split() if p.isalpha()]
    if len(parts) >= 2:
        return parts[0].title(), parts[-1].title()
    if len(parts) == 1:
        return parts[0].title(), "Candidate"
    return "Unknown", "Candidate"


def _guess_name(text: str, path: Path) -> tuple[str, str]:
    # Try the first few non-empty lines for a plausible "Firstname Lastname".
    for line in (ln.strip() for ln in text.splitlines()[:8]):
        if 0 < len(line) <= 40 and "@" not in line and not any(ch.isdigit() for ch in line):
            words = [w for w in re.split(r"\s+", line) if w.isalpha()]
            if 2 <= len(words) <= 4:
                return words[0].title(), words[-1].title()
    return _name_from_filename(path)


def _detect(text_lower: str, vocab: dict[str, list[str]]) -> str | None:
    for label, needles in vocab.items():
        for needle in needles:
            keyword = needle.strip().lower()
            if keyword and re.search(
                rf"(?<![a-z0-9]){re.escape(keyword)}(?![a-z0-9])",
                text_lower,
            ):
                return label
    return None


def _clean_text(value: str | None) -> str:
    text = str(value or "")
    text = text.replace("â€¢", " ").replace("â€“", "-").replace("â€”", "-")
    text = text.replace("\u2022", " ").replace("\u2013", "-").replace("\u2014", "-")
    text = re.sub(r"[_\-]{4,}", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip(" \t\r\n-|,;")


def _is_noisy_card_text(value: str | None) -> bool:
    text = _clean_text(value)
    if not text:
        return True
    lower = text.lower()
    if len(text) > 90:
        return True
    if EMAIL_RE.search(text) or PHONE_RE.search(text):
        return True
    if re.search(r"^\d+\s+", text):
        return True
    if any(word in lower for word in NOISY_WORDS):
        return True
    if any(re.search(rf"\b{re.escape(word)}\b", lower) for word in ADDRESS_WORDS):
        return True
    if re.search(r"[\u2022|]{1,}", text):
        return True
    return False


def _title_name(value: str | None) -> str | None:
    text = _clean_text(value)
    if not text:
        return None
    return " ".join(part[:1].upper() + part[1:].lower() for part in text.split())


def _bad_name(first: str | None, last: str | None) -> bool:
    first_text = _clean_text(first)
    last_text = _clean_text(last)
    joined = f"{first_text} {last_text}".strip().lower()
    if not first_text or not last_text:
        return True
    if any(ch.isdigit() for ch in joined):
        return True
    if joined in {"unknown candidate", "unknown provider", "new member"}:
        return True
    if first_text.lower() in SECTION_TEXT or last_text.lower() in SECTION_TEXT:
        return True
    cred_upper = {c.upper().replace(".", "") for c in CREDENTIALS}
    if first_text.upper().replace(".", "") in cred_upper or last_text.upper().replace(".", "") in cred_upper:
        return True
    return False


def _title_city(value: str | None) -> str | None:
    city = _clean_text(value)
    if not city or len(city) <= 2:
        return None
    if len(city.split()) > 4:
        return None
    lower = city.lower()
    if any(re.search(rf"\b{re.escape(word)}\b", lower) for word in CITY_BAD_WORDS):
        return None
    if any(re.search(rf"\b{re.escape(word)}\b", lower) for word in ADDRESS_WORDS):
        return None
    compact = city.upper().replace(".", "").replace(" ", "")
    if compact in {c.upper().replace(".", "") for c in CREDENTIALS} | {"MFA", "BFA", "BA", "BS"}:
        return None
    if "." in city and not re.search(r"\bSt\.", city):
        return None
    if any(ch.isdigit() for ch in city):
        return None
    if len(city) > 60:
        return None
    return " ".join(part[:1].upper() + part[1:].lower() for part in city.split())


def _clean_state(value: str | None) -> str | None:
    state = _clean_text(value).upper()
    if state in US_STATES:
        return state
    lower = _clean_text(value).lower()
    for code, name in US_STATES.items():
        if lower == name:
            return code
    return None


def _clean_profession(value: str | None) -> str | None:
    text = _clean_text(value)
    if not text or _is_noisy_card_text(text):
        return None
    upper = text.upper().replace(".", "")
    aliases = {"PA C": "PA-C", "PHARMD": "PharmD"}
    upper = aliases.get(upper, upper)
    for code in [*PROFESSION_FROM_CRED, "CRNA", "LPN", "CNA", "LVN", "RT", "PT", "OT"]:
        if upper == code.upper():
            return "PharmD" if code.upper() == "PHARMD" else code.upper()
    if len(upper) <= 12 and re.fullmatch(r"[A-Z-]+", upper):
        return upper
    return text[:50]


def _line_tokens_for_name(line: str) -> list[str]:
    cleaned = _clean_text(line).replace(",", " ")
    cleaned = re.split(r"\s+\d", cleaned, 1)[0]
    tokens = []
    for token in re.split(r"\s+", cleaned):
        letters = token.replace(".", "").replace("'", "").replace("-", "")
        if letters.isalpha():
            tokens.append(token)
    cred_upper = {c.upper().replace(".", "") for c in CREDENTIALS}
    return [t for t in tokens if t.upper().replace(".", "") not in cred_upper]


def _candidate_name_from_line(line: str) -> tuple[str, str] | None:
    explicit = re.search(r"\bname\s*:\s*([^,@|0-9]+)", line, re.IGNORECASE)
    if explicit:
        tokens = _line_tokens_for_name(explicit.group(1))
        if 2 <= len(tokens) <= 5:
            return _title_name(tokens[0]) or tokens[0], _title_name(tokens[-1]) or tokens[-1]
    if _is_noisy_card_text(line) and not re.match(r"^[^\d@|,]+?\s+\d", line):
        return None
    tokens = _line_tokens_for_name(line)
    if 2 <= len(tokens) <= 5:
        return _title_name(tokens[0]) or tokens[0], _title_name(tokens[-1]) or tokens[-1]
    lower = _clean_text(line).lower()
    if lower in SECTION_TEXT or any(h in lower for h in SECTION_TEXT):
        return None
    return None


def _guess_name(text: str, path: Path) -> tuple[str, str]:
    # Try the first few non-empty lines for a plausible "Firstname Lastname".
    for line in (ln.strip() for ln in text.splitlines()[:14]):
        name = _candidate_name_from_line(line)
        if name:
            return name
    return _name_from_filename(path)


def _extract_profession(text: str, fallback: str | None = None) -> str | None:
    profession = _clean_profession(fallback)
    if profession:
        return profession
    header = "\n".join(text.splitlines()[:18])
    for code in ("CRNA", "CNM", "PA-C", "PharmD", "LPN", "LVN", "CNA", "RN", "NP", "PA",
                 "RT", "PT", "OT", "MD", "DO", "MBBS"):
        if re.search(rf"\b{re.escape(code)}\b", header, re.IGNORECASE):
            return "PharmD" if code.lower() == "pharmd" else code.upper()
    tl = " " + text.lower() + " "
    detected = _detect(tl, PROFESSIONS)
    return _clean_profession(detected)


def _clean_specialty(value: str | None) -> str | None:
    text = _clean_text(value)
    if _is_noisy_card_text(text):
        return None
    return text[:100]


def _extract_location(text: str) -> tuple[str | None, str | None]:
    best: tuple[int, str, str] | None = None
    lines = [_clean_text(ln) for ln in text.splitlines() if _clean_text(ln)]
    for idx, line in enumerate(lines[:40]):
        for match in list(CITY_STATE_RE.finditer(line)) + list(CITY_FULL_STATE_RE.finditer(line)):
            city = _title_city(match.group(1))
            state = _clean_state(match.group(2))
            if not city or not state:
                continue
            has_location_context = (
                bool(re.search(r"\b\d{5}(?:-\d{4})?\b", line))
                or EMAIL_RE.search(line)
                or PHONE_RE.search(line)
                or "|" in line
                or any(re.search(rf"\b{re.escape(word)}\b", line.lower()) for word in ADDRESS_WORDS)
            )
            if state == "MD" and not has_location_context:
                continue
            if idx > 12 and not has_location_context:
                continue
            score = 100 - idx
            if re.search(r"\b\d{5}(?:-\d{4})?\b", line):
                score += 20
            if EMAIL_RE.search(line) or PHONE_RE.search(line) or "|" in line:
                score += 10
            if best is None or score > best[0]:
                best = (score, city, state)
    if best:
        return best[1], best[2]

    for code, name in US_STATES.items():
        if re.search(rf"\b{code}\b", text) or re.search(rf"\b{re.escape(name)}\b", text, re.IGNORECASE):
            return None, code
    return None, None


def _extract_years(text: str, fallback: int | None = None) -> int:
    value = int(fallback or 0)
    if value < 0 or value > 60:
        value = 0
    matches = [int(m.group(1)) for m in YEARS_RE.finditer(text)]
    matches = [m for m in matches if 0 < m <= 60]
    return max(matches) if matches else value


def format_resume_fields(fields: dict, text: str, path: Path) -> dict:
    """Normalize parsed resume fields for provider cards and profile filters."""
    out = dict(fields)
    first, last = _guess_name(text, path)
    if first and last and _bad_name(out.get("first_name"), out.get("last_name")):
        out["first_name"] = first[:100]
        out["last_name"] = last[:100]
    else:
        out["first_name"] = (_title_name(out.get("first_name")) or "Unknown")[:100]
        out["last_name"] = (_title_name(out.get("last_name")) or "Provider")[:100]

    email_m = EMAIL_RE.search(text)
    phone_m = PHONE_RE.search(text)
    if email_m:
        out["email"] = email_m.group(0)[:255]
    if phone_m:
        out["phone"] = phone_m.group(0)[:30]

    profession = _extract_profession(text, out.get("profession_type"))
    out["profession_type"] = profession[:50] if profession else None

    specialty = _clean_specialty(out.get("specialty")) or _detect(" " + text.lower() + " ", SPECIALTIES)
    out["specialty"] = specialty[:100] if specialty else None

    parsed_city, parsed_state = _extract_location(text)
    city = _title_city(out.get("city")) or parsed_city
    state = _clean_state(out.get("state_code")) or parsed_state
    out["city"] = city[:120] if city else None
    out["state_code"] = state

    out["years_experience"] = _extract_years(text, out.get("years_experience"))

    headline = _clean_text(out.get("headline"))
    if _is_noisy_card_text(headline):
        headline = None
    if not headline:
        bits = [out.get("specialty"), out.get("profession_type")]
        headline = " ".join(b for b in bits if b) or None
    out["headline"] = headline[:255] if headline else None

    out["american_board"] = primary_american_board(out.get("certifications")) or out.get("american_board")
    category = classify_provider(out.get("profession_type"), out.get("specialty"), out.get("headline"))
    out["provider_category"] = category
    return out


# --- Structured profile parsing (scraped physician/provider profiles) -----
# Format:  Line0: "Name CRED CRED"
#          Line1: "Specialty • City, ST"
#          Line2: focus/role (optional)
#          ... "Certifications & Licensure" section with "ST State Medical License"

SECTION_HEADERS = (
    "education & training", "certifications & licensure", "awards",
    "publications", "professional memberships", "languages", "experience",
)
LOCATION_RE = re.compile(r"(.+?)\s*[•·|‣∙·]\s*(.+?),\s*([A-Z]{2})\b")
LOCATION_FALLBACK_RE = re.compile(r"^(.+?),\s*([A-Z]{2})\b")
STATE_LICENSE_RE = re.compile(
    r"\b([A-Z]{2})\s+State Medical License", re.IGNORECASE)
LICENSE_YEARS_RE = re.compile(r"(\d{4})\s*[-–]\s*(\d{4})")


def _strip_credentials(name_line: str) -> tuple[str, str, str | None]:
    """Return (first, last, profession) from a 'Name MD, MPH' style line."""
    cleaned = name_line.replace(",", " ")
    tokens = [t for t in re.split(r"\s+", cleaned) if t]
    profession = None
    for cred in PROFESSION_FROM_CRED:
        if any(t.upper().strip(".") == cred for t in tokens):
            profession = cred
            break
    cred_upper = {c.upper() for c in CREDENTIALS}
    name_tokens = [
        t for t in tokens
        if t.upper().strip(".") not in cred_upper and "(" not in t and ")" not in t
    ]
    if not name_tokens:
        name_tokens = tokens
    first = name_tokens[0].title() if name_tokens else "Unknown"
    last = name_tokens[-1].title() if len(name_tokens) > 1 else "Provider"
    return first, last, profession


def _looks_structured(lines: list[str]) -> bool:
    head = "\n".join(lines[:6]).lower()
    return ("certifications & licensure" in "\n".join(lines).lower()
            or "education & training" in head
            or bool(lines and LOCATION_RE.search(lines[1] if len(lines) > 1 else "")))


def parse_structured(text: str, path: Path) -> dict:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    first, last, profession = _strip_credentials(lines[0]) if lines else ("Unknown", "Provider", None)

    specialty = city = state_code = None
    headline = None
    # Scan the first few lines for the "Specialty • City, ST" header — the name
    # can wrap across lines, so its position isn't fixed.
    loc_idx = None
    for i in range(1, min(len(lines), 6)):
        m = LOCATION_RE.search(lines[i])
        if m:
            specialty = m.group(1).strip()
            city = m.group(2).strip()
            state_code = m.group(3).upper()
            loc_idx = i
            break
    if loc_idx is None:  # fallback: a bare "City, ST" line
        for i in range(1, min(len(lines), 6)):
            fm = LOCATION_FALLBACK_RE.search(lines[i])
            if fm and not any(h in lines[i].lower() for h in SECTION_HEADERS):
                city, state_code = fm.group(1).strip(), fm.group(2).upper()
                loc_idx = i
                break

    # The line after the location (focus / role) becomes the headline.
    if loc_idx is not None and loc_idx + 1 < len(lines):
        nxt = lines[loc_idx + 1]
        if not any(h in nxt.lower() for h in SECTION_HEADERS):
            headline = nxt[:250]
    if not headline:
        headline = specialty

    # Licenses: "NC State Medical License" followed by a "1993 - 2026" line.
    licenses = []
    seen_states = set()
    for i, ln in enumerate(lines):
        lm = STATE_LICENSE_RE.search(ln)
        if lm:
            st = lm.group(1).upper()
            if st in seen_states:
                continue
            seen_states.add(st)
            expiry_year = None
            for nxt in lines[i + 1: i + 3]:
                ym = LICENSE_YEARS_RE.search(nxt)
                if ym:
                    expiry_year = int(ym.group(2))
                    break
            licenses.append({
                "state_code": st,
                "license_type": profession or "MD",
                "expiry_year": expiry_year,
            })
    if state_code is None and licenses:
        state_code = licenses[0]["state_code"]

    # Board certifications → certification names.
    board_certs = [
        ln.strip() for ln in lines
        if ln.lower().startswith("american board of")
        or ln.lower().startswith("abms")
    ]
    board_certs = list(dict.fromkeys(board_certs))[:6]

    npi_m = NPI_RE.search(text)
    email_m = EMAIL_RE.search(text)

    # A short bio from the role + first education line.
    bio_parts = []
    if headline:
        bio_parts.append(headline)
    for i, ln in enumerate(lines):
        if ln.lower().startswith("education & training") and i + 1 < len(lines):
            bio_parts.append(f"Trained at {lines[i + 1]}")
            break
    bio = " · ".join(bio_parts) or None

    prof = profession or "MD"
    fields = {
        "first_name": first,
        "last_name": last,
        "email": email_m.group(0) if email_m else None,
        "phone": None,
        "profession_type": prof,
        "specialty": specialty,
        "city": city,
        "state_code": state_code,
        "certifications": board_certs,
        "licenses": licenses,
        "years_experience": 0,
        "npi_number": npi_m.group(1) if npi_m else None,
        "headline": headline,
        "bio": bio,
        "american_board": primary_american_board(board_certs),
        "provider_category": classify_provider(prof, specialty, headline),
    }
    return format_resume_fields(fields, text, path)


def parse_resume(text: str, path: Path) -> dict:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if _looks_structured(lines):
        return parse_structured(text, path)

    tl = " " + text.lower() + " "
    first, last = _guess_name(text, path)

    email_m = EMAIL_RE.search(text)
    phone_m = PHONE_RE.search(text)
    years_m = YEARS_RE.search(text)
    npi_m = NPI_RE.search(text)

    profession = _detect(tl, PROFESSIONS)
    if profession is None:
        # Fallback: a standalone credential token (e.g. "...Department RN").
        for code in ("CRNA", "CNM", "PharmD", "LPN", "CNA", "RN", "NP", "PA",
                     "RT", "PT", "OT", "MD", "DO", "MBChB"):
            if re.search(rf"\b{code}\b", text):
                profession = code
                break
    specialty = _detect(tl, SPECIALTIES)

    certs = [c for c in CERTIFICATIONS if re.search(rf"\b{c}\b", text, re.IGNORECASE)]

    state_code = None
    # Prefer explicit 2-letter codes near "RN" / state lines, else full names.
    for code, name in US_STATES.items():
        if re.search(rf"\b{code}\b", text) or name in tl:
            state_code = code
            break

    headline_bits = [b for b in [specialty, profession] if b]
    headline = " ".join(headline_bits) if headline_bits else None
    if headline and years_m:
        headline += f" · {years_m.group(1)} yrs"

    fields = {
        "first_name": first,
        "last_name": last,
        "email": email_m.group(0) if email_m else None,
        "phone": phone_m.group(0) if phone_m else None,
        "profession_type": profession,
        "specialty": specialty,
        "city": None,
        "certifications": certs,
        "licenses": [],
        "years_experience": int(years_m.group(1)) if years_m else 0,
        "state_code": state_code,
        "npi_number": npi_m.group(1) if npi_m else None,
        "headline": headline,
        "bio": None,
        "american_board": primary_american_board(certs),
        "provider_category": classify_provider(profession, specialty, headline),
    }
    return format_resume_fields(fields, text, path)
